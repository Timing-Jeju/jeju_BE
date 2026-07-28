from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/business/timing-jeju-business-registration-plan.docx")

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9D9D9"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def border_cell(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", 90), ("bottom", 90), ("start", 140), ("end", 140)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border_cell(cell)
            margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run(run, size=10)
        if r_idx == 0:
            for cell in row.cells:
                shade(cell, HEADER_FILL)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        set_run(run, size=10, bold=True, color=DARK)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        set_run(run, size=16, bold=True, color=BLUE)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run(run, size=13, bold=True, color=BLUE)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run(run, size=11, color=INK)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        run = p.add_run(item)
        set_run(run, size=11, color=INK)


def callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    border_cell(cell)
    margins(cell)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_run(run, size=11, bold=True, color=DARK)
    p.add_run("\n")
    run = p.add_run(text)
    set_run(run, size=10.5, color=INK)
    p.paragraph_format.line_spacing = 1.2
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "내용"
    for key, value in rows:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
    style_table(table)
    return table


def matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Inches(w)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    style_table(table)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("타이밍제주 사업자등록용 사업계획서")
    set_run(run, size=20, bold=True, color=RGBColor(11, 37, 69))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("제주 뚜벅이 일정 운영 앱·웹 서비스")
    set_run(run, size=12, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    callout(
        doc,
        "작성 목적",
        "본 문서는 사업자등록 신청 및 초기 사업 설명을 위한 사업계획서 초안이다. GPS 활용에 따른 위치기반서비스사업 신고는 사업자등록과 별도로 검토하며, 본 문서에서는 향후 인허가·신고 확인사항으로 정리한다.",
    )

    heading(doc, "1. 사업 개요")
    kv_table(doc, [
        ("상호", "타이밍제주 또는 추후 확정 상호"),
        ("사업 형태", "개인사업자 또는 법인사업자 등록 검토"),
        ("업태 후보", "정보통신업, 서비스업"),
        ("종목 후보", "응용 소프트웨어 개발 및 공급업, 모바일 앱 개발, 온라인 정보 제공업"),
        ("사업 내용", "제주 뚜벅이 여행자를 위한 일정 생성, 대중교통 기반 이동 타이밍 안내, 여행 데이터 서비스"),
        ("서비스 채널", "모바일 앱, 웹 서비스, 향후 관리자·지역관광 데이터 화면"),
        ("개업 예정일", "사업자등록 신청 시 확정"),
        ("사업장 소재지", "사업자등록 신청 시 임대차계약서 또는 자가 사업장 기준으로 기재"),
    ])
    body(
        doc,
        "타이밍제주는 제주를 렌터카 없이 여행하는 이용자가 버스 시간 때문에 일정을 실패하지 않도록 돕는 앱·웹 서비스이다. 사용자는 가고 싶은 관광지와 여행 시간을 입력하거나 직접 일정을 구성할 수 있고, 서비스는 관광지 정보와 버스 데이터를 바탕으로 일정 실행 가능성을 안내한다.",
    )

    heading(doc, "2. 창업 배경 및 필요성")
    body(
        doc,
        "제주도는 관광 수요가 높은 지역이지만, 렌터카 없이 이동하는 관광객에게는 일정 운영 난이도가 높다. 관광지 사이의 거리가 길고, 정류장 접근 시간과 버스 배차 간격에 따라 하루 일정이 크게 흔들릴 수 있다. 기존 지도 서비스는 구간 길찾기에 강하지만 여러 관광지를 방문하는 하루 일정 전체가 버스 시간에 맞게 유지되는지까지 관리해주지는 않는다.",
    )
    body(
        doc,
        "타이밍제주는 이 문제를 관광지 추천이 아니라 일정 운영 문제로 본다. 이용자가 언제 출발해야 하는지, 버스를 놓치면 얼마나 기다려야 하는지, 일정이 밀렸을 때 무엇을 줄이거나 대체해야 하는지를 알려주는 것이 서비스의 핵심이다.",
    )

    heading(doc, "3. 주요 서비스")
    bullets(doc, [
        "자연어 일정 입력: 사용자가 문장으로 원하는 장소와 시간을 입력하면 장소와 시간 조건을 구조화한다.",
        "직접 일정 만들기: 사용자가 날짜, 출발지, 복귀 지점, 방문 장소, 체류시간, 장소 순서를 직접 설정한다.",
        "관광지 후보 매칭: TourAPI 관광정보를 활용해 입력한 장소명을 실제 관광지 후보와 연결한다.",
        "버스 기반 실행 가능성 검사: 정류장, 노선, 시간표, 실시간 도착정보를 바탕으로 출발 권장 시각과 놓침 위험을 계산한다.",
        "일정 안전도 안내: 전체 일정과 구간별 위험 사유를 점수 또는 상태로 보여준다.",
        "놓침 복구안 제공: 버스를 놓쳤거나 일정이 밀렸을 때 다음 버스, 체류시간 단축, 장소 제외, 대체 후보를 제안한다.",
        "AI 설명: AI는 서버 계산 결과를 이용자가 이해하기 쉬운 문장으로 설명하며, 버스 시간이나 노선 사실을 임의로 생성하지 않는다.",
    ])

    heading(doc, "4. 목표 고객")
    matrix(doc, ["고객군", "니즈", "제공 가치"], [
        ("제주 뚜벅이 여행자", "버스 중심 일정 설계와 이동 타이밍 판단", "출발 권장 시각, 일정 안전도, 복구안 제공"),
        ("제주 초행 관광객", "관광지와 정류장 위치 파악", "TourAPI 관광지 후보, 가까운 정류장 후보 안내"),
        ("청년·대학생 여행자", "렌터카 없이 저비용 여행", "대중교통 기반 실현 가능한 일정 제공"),
        ("외국인 관광객", "한국어 버스 정보 이해 어려움", "향후 다국어 안내와 직관적 타임라인 제공"),
        ("지역관광기관", "뚜벅이 이동 불편 구간 파악", "익명·집계 기반 RTO 분석 확장 가능"),
    ], widths=[1.7, 2.8, 3.0])

    heading(doc, "5. 수익 모델")
    bullets(doc, [
        "초기 무료 서비스: 공모전 및 베타 테스트 단계에서 사용자 반응과 핵심 기능을 검증한다.",
        "프리미엄 기능: 실시간 버스 보정, 다일차 일정 저장, 고도화된 알림, 다국어 안내 등을 유료 기능으로 검토한다.",
        "지역 제휴: 짐 보관소, 숙소, 카페, 체험 시설, 관광지 입장권 등과 연계한 수수료 모델을 검토한다.",
        "B2B 데이터 리포트: 이동 위험 구간, 대기 시간, 관광권역별 뚜벅이 난이도 등을 익명·집계 데이터로 제공한다.",
        "공공·관광 사업 연계: 지역관광 활성화, 대중교통 관광 접근성 개선 사업과 연계한다.",
    ])

    heading(doc, "6. 개발 및 사업화 계획")
    matrix(doc, ["단계", "개발 범위", "사업화 목표"], [
        ("본선 MVP", "제주 동쪽 대표 코스, 시간표 기반 계산, 직접 일정 만들기, 복구안", "공모전 본선에서 작동하는 서비스 증명"),
        ("실사용 MVP", "실시간 버스 도착·위치 연동, 알림 보정, GPS 기반 현재 상태 안내", "실제 여행 중 신뢰 가능한 앱으로 검증"),
        ("제주 전역 확장", "더 많은 정류장·노선·시간표, 숙소권, 관광권역 확대", "서비스 사용 범위 확대"),
        ("사업 고도화", "날씨 보정, 짐 보관, 택시 플랜B, 다국어, RTO 대시보드", "수익 모델과 지역관광 데이터 가치 확보"),
    ], widths=[1.4, 3.4, 2.7])

    heading(doc, "7. 위치정보 및 인허가 검토")
    body(
        doc,
        "사업자등록의 주된 목적은 소프트웨어·앱 서비스 개발 및 운영 사업을 개시하는 것이다. 다만 타이밍제주는 향후 이용자 GPS 위치를 활용해 현재 위치 기반 일정 상태, 가까운 정류장 후보, 출발 알림을 제공할 예정이므로 위치기반서비스사업 신고 여부를 별도로 확인해야 한다.",
    )
    matrix(doc, ["구분", "사업자등록 문서 내 처리", "추가 확인"], [
        ("GPS 위치 활용", "서비스 기능 설명에 포함", "위치정보 이용 동의, 보유기간, 동의 철회 방식 설계"),
        ("위치기반서비스 신고", "사업자등록과 별도 확인사항으로 기재", "방송미디어통신위원회·위치정보지원센터 양식 확인"),
        ("개인정보처리방침", "운영 준비사항으로 기재", "서비스 출시 전 약관과 처리방침 작성"),
        ("통신판매", "초기 무료 서비스라면 우선 제외 가능", "유료 결제나 상품 중개 시 별도 신고 검토"),
    ], widths=[1.6, 3.1, 2.8])

    heading(doc, "8. 팀 구성")
    matrix(doc, ["역할", "주요 담당"], [
        ("백엔드 1", "TourAPI, 제주 정류장·노선·시간표, 실시간 버스 API, 데이터 헬스, DB 구조"),
        ("백엔드 2", "일정 계산 엔진, 출발 권장 시각, 위험도, 복구안, AI 파싱·설명 가드레일"),
        ("프론트엔드", "자연어 입력, 직접 일정 만들기, 결과 화면, 타임라인, 알림, 지도 딥링크"),
        ("디자이너 1", "입력 흐름, 직접 일정 빌더, 장소 추가·삭제·순서 조정, 복구 UX"),
        ("디자이너 2", "위험 카드, 알림 문구, 불안감 낮은 UX 라이팅, 발표용 화면"),
    ], widths=[1.4, 6.0])

    heading(doc, "9. 운영 계획")
    bullets(doc, [
        "초기에는 공모전 본선과 팀 내부 테스트를 통해 대표 코스의 기능 안정성을 검증한다.",
        "베타 테스트는 제주 뚜벅이 여행자와 지인 테스트 그룹을 대상으로 진행한다.",
        "외부 API 장애에 대비해 fixture 데이터와 fallback 문구를 준비한다.",
        "사용자 위치정보는 서비스 제공 목적에 필요한 범위에서만 사용하고 장기 보관을 지양한다.",
        "정식 출시 전 위치정보 이용약관, 개인정보처리방침, 위치정보관리책임자, 민원 대응 절차를 확정한다.",
    ])

    heading(doc, "10. 예상 비용")
    matrix(doc, ["항목", "초기 비용 방향", "비고"], [
        ("서버/호스팅", "무료 또는 저가 클라우드", "공모전 MVP는 로컬·저비용 운영 가능"),
        ("데이터베이스", "무료 또는 저가 PostgreSQL", "PostGIS 사용 가능 환경 검토"),
        ("OpenAI API", "테스트·소량 호출 기준 저비용", "실시간 상태 계산에는 사용하지 않음"),
        ("지도/API", "무료 사용량 우선", "사용량 증가 시 비용 검토"),
        ("도메인/배포", "소액", "서비스명 확정 후 구매 검토"),
        ("법무/신고", "필요 시 상담 비용", "위치기반서비스 신고, 약관 검토"),
    ], widths=[1.5, 2.6, 3.4])

    heading(doc, "11. 사업자등록 시 확인할 사항")
    bullets(doc, [
        "개인사업자 또는 법인사업자 중 어떤 형태로 등록할지 결정한다.",
        "상호, 사업장 주소, 개업일, 주업종과 부업종을 확정한다.",
        "주업종은 소프트웨어 개발 및 공급업 또는 온라인 정보 제공업 계열로 검토한다.",
        "유료 결제, 광고, 제휴 판매를 시작할 경우 통신판매업 신고 여부를 별도로 확인한다.",
        "GPS 기반 기능을 실제 서비스에 제공하기 전 위치기반서비스사업 신고 여부를 확인한다.",
        "사업자등록 신청은 홈택스 또는 관할 세무서를 통해 진행한다.",
    ])

    heading(doc, "12. 참고")
    bullets(doc, [
        "국세청 안내에 따르면 홈택스를 통한 사업자등록 신청과 구비서류 전자제출이 가능하다.",
        "위치기반서비스사업 관련 신고는 방송미디어통신위원회 안내와 위치정보지원센터 자료를 기준으로 별도 준비한다.",
        "본 문서는 사업자등록과 초기 사업 설명을 위한 초안이며, 실제 신청 전 세무서·전문가 확인을 권장한다.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build().resolve())
