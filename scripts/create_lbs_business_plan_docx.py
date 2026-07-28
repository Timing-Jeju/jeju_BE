from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/business/timing-jeju-lbs-business-plan.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "D9D9D9"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=11):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                set_paragraph_font(paragraph, 10)
        if header and row_idx == 0:
            for cell in row.cells:
                shade_cell(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_font(run, size=10, bold=True, color=DARK_BLUE)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.333
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        run = paragraph.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.208
        run = paragraph.add_run(item)
        set_run_font(run, size=11)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    set_cell_border(cell, "C9D3DF")
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    set_run_font(run, size=11, bold=True, color=DARK_BLUE)
    paragraph.add_run("\n")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    paragraph.paragraph_format.line_spacing = 1.2
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.8)
    table.rows[0].cells[0].text = "항목"
    table.rows[0].cells[1].text = "내용"
    for key, value in rows:
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = value
    style_table(table)
    return table


def add_matrix_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    style_table(table)
    return table


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("타이밍제주 위치기반서비스사업 사업계획서 초안")
    set_run_font(run, size=9, color=RGBColor(90, 90, 90))


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("타이밍제주 위치기반서비스사업 사업계획서")
    set_run_font(run, size=20, bold=True, color=RGBColor(11, 37, 69))
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("GPS 기반 제주 뚜벅이 일정 운영 서비스")
    set_run_font(run, size=12, color=RGBColor(90, 90, 90))
    subtitle.paragraph_format.space_after = Pt(18)

    add_callout(
        doc,
        "작성 목적",
        "본 문서는 타이밍제주 서비스가 이용자 GPS 위치를 활용하기 위해 위치기반서비스사업 신고를 준비하는 사업계획서 초안이다. 실제 제출 전에는 방송미디어통신위원회 안내 양식, 위치정보지원센터 상담, 세무·법무 검토를 통해 사업자 정보와 보호조치 내용을 확정한다.",
    )

    add_heading(doc, "1. 사업자 및 서비스 개요")
    add_kv_table(doc, [
        ("서비스명", "타이밍제주"),
        ("서비스 형태", "모바일 앱 및 웹 기반 제주 뚜벅이 일정 운영 서비스"),
        ("신고 목적", "이용자 GPS 위치를 활용한 현재 위치 기반 일정 상태 판단, 정류장 후보 안내, 출발 알림 제공"),
        ("예상 신고 유형", "위치기반서비스사업 신고 또는 소상공인등의 위치기반서비스사업 신고 대상 여부 검토"),
        ("주요 고객", "제주를 렌터카 없이 여행하는 뚜벅이 관광객"),
        ("주요 데이터", "이용자 현재 위치, TourAPI 관광정보, 제주 버스 정류장·노선·시간표, 실시간 버스 도착정보"),
    ])

    add_body(
        doc,
        "타이밍제주는 렌터카 없이 제주를 여행하는 이용자가 버스 시간 때문에 일정을 실패하지 않도록 돕는 서비스이다. 이용자는 자연어로 일정을 입력하거나 직접 장소와 체류시간을 선택해 일정을 만들 수 있고, 서비스는 현재 위치와 버스 이동 정보를 기준으로 언제 출발해야 하는지와 놓쳤을 때 어떤 복구안이 가능한지를 안내한다.",
    )

    add_heading(doc, "2. 위치기반서비스 신고 필요성")
    add_body(
        doc,
        "타이밍제주는 이용자 단말의 GPS 위치를 활용해 현재 관광지 또는 정류장 근처에 있는지 판단하고, 다음 버스 탑승 가능성, 출발 권장 시각, 남은 체류 가능 시간, 놓침 위험을 안내한다. 따라서 사용자의 위치정보를 서비스 기능에 활용하는 구조이며, 위치기반서비스사업 신고 검토가 필요하다.",
    )
    add_body(
        doc,
        "본 서비스는 자체적으로 통신망이나 측위 설비를 운영해 위치정보를 다른 사업자에게 제공하는 위치정보사업을 목표로 하지 않는다. 이용자의 동의를 받아 앱 서비스 제공 목적 범위 안에서 위치를 활용하는 위치기반서비스 형태로 설계한다.",
    )

    add_heading(doc, "3. 위치정보 이용 목적")
    add_matrix_table(
        doc,
        ["이용 목적", "서비스 내 활용", "이용자 가치"],
        [
            ("현재 위치 확인", "현재 장소 또는 정류장 주변 여부 판단", "여행 중 현재 일정 상태를 빠르게 확인"),
            ("가까운 정류장 후보 안내", "관광지 주변 탑승·하차 정류장 후보 검색", "정류장 탐색 시간 감소"),
            ("출발 권장 시각 계산", "정류장까지 도보 시간과 버스 시간표를 결합", "버스를 놓칠 위험 감소"),
            ("라이브 일정 상태 안내", "현재 위치와 현재 시각을 기준으로 여유·출발준비·위험 상태 표시", "지금 떠날지 머물지 판단 가능"),
            ("알림 제공", "출발 권장 시각 전 로컬 또는 앱 알림 제공", "여행 중 앱을 계속 보지 않아도 일정 관리 가능"),
            ("놓침 복구안 제공", "버스를 놓쳤거나 일정이 밀린 경우 다음 버스와 대체 일정 계산", "하루 일정 붕괴 방지"),
        ],
        widths=[1.6, 3.1, 2.8],
    )

    add_heading(doc, "4. 수집 및 이용하는 위치정보")
    add_matrix_table(
        doc,
        ["구분", "수집·이용 항목", "보유 및 처리 원칙"],
        [
            ("현재 위치", "위도, 경도, 위치 확인 시각, 위치 정확도", "일정 상태 계산에 사용하고 장기 보관하지 않음"),
            ("수동 대체 위치", "이용자가 선택한 가까운 정류장 또는 관광지", "GPS 권한 거절 시 대체 기준으로 사용"),
            ("일정 위치", "출발지, 방문 관광지, 복귀 지점 좌표", "이용자가 만든 일정 계산을 위해 사용"),
            ("알림 기준 위치", "다음 이동 구간의 관광지·정류장 위치", "출발 알림과 이동 안내 기준으로 사용"),
            ("로그 및 오류", "위치정보 원문이 아닌 오류 코드, fallback 여부, 데이터 기준 시각", "서비스 품질 개선 및 장애 대응에 사용"),
        ],
        widths=[1.35, 3.2, 2.95],
    )
    add_body(
        doc,
        "정확한 GPS 좌표는 서비스 제공에 필요한 범위에서만 사용한다. 이용자가 위치 권한을 거절하는 경우에도 가까운 정류장이나 관광지를 직접 선택해 기능을 사용할 수 있도록 대체 흐름을 제공한다.",
    )

    add_heading(doc, "5. 서비스 제공 흐름")
    add_numbered(doc, [
        "이용자가 위치정보 이용에 동의하거나 현재 위치 대신 정류장·관광지를 직접 선택한다.",
        "이용자가 자연어로 일정을 입력하거나 직접 일정 만들기 화면에서 방문 장소와 체류시간을 입력한다.",
        "서비스는 TourAPI를 통해 관광지 후보를 매칭하고, 관광지 좌표 기준 가까운 정류장 후보를 조회한다.",
        "버스 노선과 시간표, 실시간 도착정보를 기준으로 이동 가능성, 출발 권장 시각, 놓침 위험을 계산한다.",
        "현재 위치와 현재 시각을 기준으로 여유, 출발 준비, 위험 상태를 표시한다.",
        "출발 권장 시각 전 알림을 제공하고, 버스를 놓친 경우 복구안을 계산한다.",
        "AI는 서버가 계산한 사실만 바탕으로 위험 이유와 조정 제안을 설명한다.",
    ])

    add_heading(doc, "6. 위치정보 보호조치")
    add_matrix_table(
        doc,
        ["보호 영역", "적용 계획", "비고"],
        [
            ("동의", "위치정보 이용 전 명확한 동의 화면 제공", "목적, 항목, 보유기간, 철회 방법 표시"),
            ("최소 수집", "정확한 GPS가 불필요한 경우 정류장·관광지 선택으로 대체", "위치 권한 거절 흐름 제공"),
            ("목적 제한", "일정 계산, 정류장 안내, 출발 알림, 복구안 제공에 한정", "마케팅 목적 별도 동의 없이 사용하지 않음"),
            ("보관 제한", "현재 위치 원문은 장기 저장하지 않고 필요한 경우 익명·집계 통계만 사용", "구체 보관기간은 약관에서 확정"),
            ("접근 통제", "서버 API 접근 권한, 관리자 접근 로그, API key 서버 보관", "프론트에 OpenAI/API key 노출 금지"),
            ("암호화", "전송 구간 HTTPS 적용, 민감 설정값 환경변수 관리", "운영 전 보안 점검 필요"),
            ("제3자 제공", "이용자 위치정보를 제3자에게 판매하거나 제공하지 않음", "필요 시 별도 동의 후 처리"),
            ("파기", "목적 달성 또는 이용자 요청 시 위치 관련 저장정보 파기", "파기 절차 문서화"),
        ],
        widths=[1.4, 4.3, 2.1],
    )

    add_heading(doc, "7. 이용자 권리 및 동의 철회")
    add_bullets(doc, [
        "이용자는 위치정보 이용 동의를 하지 않아도 수동 위치 선택 방식으로 일부 기능을 이용할 수 있다.",
        "이용자는 앱 설정 또는 서비스 설정 화면에서 위치정보 이용 동의를 철회할 수 있다.",
        "이용자는 본인의 위치정보 이용 내역 또는 처리 기준에 대한 문의를 할 수 있다.",
        "위치정보 이용 동의 철회 시 현재 위치 기반 상태 안내와 자동 알림 일부 기능은 제한될 수 있다.",
        "미성년자 또는 보호가 필요한 이용자 관련 처리는 실제 서비스 출시 전 별도 정책을 검토한다.",
    ])

    add_heading(doc, "8. 주요 설비 및 시스템 구성")
    add_matrix_table(
        doc,
        ["구성 요소", "역할", "설치·운영 방식"],
        [
            ("모바일 앱", "위치 권한 요청, 현재 위치 확인, 일정 입력, 알림 표시", "Expo 기반 앱, 이용자 단말에서 실행"),
            ("백엔드 API", "일정 계산, 위치기반 상태 판단, 복구안 생성", "클라우드 서버 또는 관리형 호스팅"),
            ("데이터베이스", "관광지, 정류장, 노선, 시간표, 일정 계산 결과 저장", "PostgreSQL/PostGIS 사용 예정"),
            ("외부 API", "TourAPI 관광정보, 제주 버스 데이터, 지도 딥링크, OpenAI 설명", "서버에서 호출하고 key는 환경변수로 관리"),
            ("데이터 헬스", "데이터 기준일, fallback, cache 상태 확인", "관리자·발표자 확인용 화면"),
        ],
        widths=[1.5, 3.4, 3.0],
    )

    add_heading(doc, "9. 개발 및 운영 단계")
    add_matrix_table(
        doc,
        ["단계", "범위", "위치정보 관련 구현"],
        [
            ("본선 MVP", "제주 동쪽 대표 코스, 시간표 기반 일정 계산", "GPS 또는 수동 위치 기준 현재 상태 판단, 정류장 후보 안내"),
            ("실사용 MVP", "실시간 버스 도착·위치 연동, 알림 보정", "실시간 위치와 버스 도착정보를 결합해 출발 알림 정확도 향상"),
            ("확장", "제주 전역, 날씨 보정, 짐 보관, 다국어, RTO 대시보드", "익명·집계 기반 관광 이동 난이도 분석"),
        ],
        widths=[1.3, 3.4, 3.2],
    )

    add_heading(doc, "10. 수익 모델 및 사업화 계획")
    add_body(
        doc,
        "초기에는 무료 서비스로 사용자 반응과 기능 완성도를 검증한다. 이후 실시간 버스 보정, 다일차 일정 저장, 고도화된 알림, 다국어 안내 등 프리미엄 기능을 검토한다. 또한 짐 보관소, 숙소, 카페, 체험 시설과의 지역 제휴 모델, 익명·집계 데이터 기반 RTO 리포트 모델을 검토할 수 있다.",
    )

    add_heading(doc, "11. 팀 구성 및 역할")
    add_matrix_table(
        doc,
        ["역할", "담당 업무"],
        [
            ("백엔드 1", "TourAPI, 제주 정류장·노선·시간표, 실시간 버스 API, 데이터 헬스, DB 구조"),
            ("백엔드 2", "일정 계산 엔진, 출발 권장 시각, 위험도, 복구안, AI 파싱·설명 가드레일"),
            ("프론트엔드", "위치 권한, 직접 일정 만들기, 결과 화면, 타임라인, 알림, 지도 딥링크"),
            ("디자이너 1", "자연어 입력과 직접 일정 만들기 UX, 위치 권한 거절 대체 흐름, 복구 흐름"),
            ("디자이너 2", "위험 상태 UI, 출발 알림 문구, 불안감 낮은 안내 문구, 발표용 화면"),
        ],
        widths=[1.5, 6.3],
    )

    add_heading(doc, "12. 제출 전 확인 필요 사항")
    add_bullets(doc, [
        "위치기반서비스사업 일반 신고와 소상공인등의 위치기반서비스사업 신고 중 어느 절차에 해당하는지 확인한다.",
        "방송미디어통신위원회 또는 전자민원센터의 최신 신고서와 사업계획서 양식을 기준으로 항목을 맞춘다.",
        "위치정보 이용약관, 개인정보처리방침, 위치정보관리책임자, 민원 대응 절차를 별도 문서로 준비한다.",
        "GPS 위치정보를 서버에 저장할지, 앱 내부 계산에만 사용할지 최종 아키텍처를 확정한다.",
        "실시간 버스 도착정보와 지도 API 이용약관상 위치정보 결합 사용에 제한이 없는지 확인한다.",
    ])

    add_heading(doc, "13. 참고한 공식 자료")
    add_bullets(doc, [
        "방송미디어통신위원회, 위치정보사업 허가 및 위치기반서비스사업 신고 안내",
        "위치정보지원센터, 위치기반서비스사업자 신고 및 소상공인등의 위치기반서비스사업자 신고 안내",
        "국가법령정보센터, 위치정보의 보호 및 이용 등에 관한 법률 및 시행령",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc().resolve())
